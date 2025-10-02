#!/usr/bin/env python3

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json
import urllib.parse
import os
import glob
import subprocess
import sys

def extract_image_filename(file_url):
    """Extract filename from Notion file URL"""
    try:
        if file_url.startswith('{{file://'):
            encoded_part = file_url[9:-2]
            decoded = urllib.parse.unquote(encoded_part)
            file_info = json.loads(decoded)

            source = file_info.get('source', '')
            if source.startswith('attachment:'):
                parts = source.split(':')
                filename = parts[2] if len(parts) > 2 else 'image.png'
                return filename
    except Exception as e:
        print(f"Error extracting filename: {e}")
    return None

def find_matching_image(filename, snapshot_folder):
    """Find matching image file in snapshot folder using exact filename matching only"""
    if not os.path.exists(snapshot_folder):
        print(f"❌ Snapshot folder not found: {snapshot_folder}")
        return None

    # Get all files in snapshot folder
    try:
        files = os.listdir(snapshot_folder)
    except Exception as e:
        print(f"❌ Error reading snapshot folder: {e}")
        return None

    # Try exact filename match first
    if filename in files:
        print(f"✅ Found exact match: {filename}")
        return os.path.join(snapshot_folder, filename)

    # Try with underscores converted to spaces
    spaced_filename = filename.replace('_', ' ')
    if spaced_filename in files:
        print(f"✅ Found match with spaces: {spaced_filename}")
        return os.path.join(snapshot_folder, spaced_filename)

    # Try matching by searching through all files (handles Unicode differences)
    expected_pattern = spaced_filename.replace('.png', '').replace('.jpg', '').replace('.jpeg', '')
    for file in files:
        if file.startswith('Screenshot') and expected_pattern in file.replace('\u202f', ' '):
            print(f"✅ Found fuzzy match: {file}")
            return os.path.join(snapshot_folder, file)

    print(f"❌ No match found for: {filename}")
    return None

def create_newsletter_with_local_images(snapshot_folder="/Users/jcastillo/Desktop/snapshots"):
    """Create newsletter using images from local snapshot folder"""

    print(f"🚀 Creating newsletter with images from: {snapshot_folder}")

    # Database items with their Notion file URLs
    items = [
        {
            'title': '🔍 @ Functionality for Cross-System Search',
            'new': 'Added @ functionality for cross-system search capabilities',
            'impact': 'Improves performance by enabling a wider range of available contexts for Distillery Chat',
            'steps': [
                'Type @ in search to access cross-system functionality',
                'Search across different system components seamlessly',
                'Get broader context for your Distillery Chat interactions'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3Aacdc2011-73f7-4e4c-971a-34c6eefe89bb%3AScreenshot_2025-10-01_at_12.52.41_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222732c038-ca37-8080-8832-f1a97980a4c2%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}',
                '{{file://%7B%22source%22%3A%22attachment%3Ae4725358-40f2-45ff-a0b0-b94b199e5666%3AScreenshot_2025-10-01_at_12.53.46_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222732c038-ca37-8080-8832-f1a97980a4c2%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        },
        {
            'title': '👁️ Enhanced GenEdit Navigation',
            'new': 'Added navigation eye icon to GenEdit tool cards to see where edits were made',
            'impact': 'Improves UI/UX of GenEdit by enabling users to focus on edited system components',
            'steps': [
                'Look for the eye icon on GenEdit tool cards',
                'Click to navigate directly to specific system components',
                'Focus on exactly what you\'re editing without distraction'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3A6f4ec752-bbbc-40e2-acf5-91c89b2635bf%3AScreenshot_2025-10-01_at_12.56.01_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222732c038-ca37-8029-8e66-ce9bb15e5122%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        },
        {
            'title': '🧪 Improvements to Regression Testing',
            'new': 'Enhanced regression testing processes for system changes',
            'impact': 'Promotes regression testing by surfacing eval test runs on the proposal page for routine edits',
            'steps': [
                'Create or edit a proposal',
                'View eval test runs directly on the proposal page',
                'Ensure comprehensive testing before implementation'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3A63b1af6e-0191-4a65-8ca0-42eb54b9548b%3AScreenshot_2025-10-01_at_1.52.48_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222722c038-ca37-80e3-9c5e-decca09533a6%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}',
                '{{file://%7B%22source%22%3A%22attachment%3Af144afb9-18ae-47ea-a9c1-b419b34aa6d6%3AScreenshot_2025-10-01_at_1.53.03_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222722c038-ca37-80e3-9c5e-decca09533a6%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        },
        {
            'title': '✏️ Improved UX of Eval Editing',
            'new': 'Enhanced user experience for editing and managing evaluations',
            'impact': 'Improves the UX of editing evals by enabling deletion and providing state indicators (edited, deleted, new)',
            'steps': [
                'Navigate to evaluation settings',
                'Use improved interface with state indicators',
                'Edit or delete evaluations more efficiently'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3Ae1e95d54-9d7c-48b9-ad73-ddaca41da02c%3AScreenshot_2025-10-01_at_5.19.21_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222702c038-ca37-8009-b31d-c49ccc29083f%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}',
                '{{file://%7B%22source%22%3A%22attachment%3A69dcba06-8323-4b56-96e2-7fe3b3bb6997%3AScreenshot_2025-10-01_at_5.14.50_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222702c038-ca37-8009-b31d-c49ccc29083f%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        },
        {
            'title': '📊 Sidebar for Eval Metadata',
            'new': 'New sidebar configuration for evaluation metadata management',
            'impact': 'Improves usability by providing a convenient in-context configuration layout for evals',
            'steps': [
                'Open evaluation settings',
                'Use the sidebar for metadata configuration',
                'Manage eval parameters without leaving your workflow'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3A39922d1b-b610-4598-bf13-31e33cc28394%3AScreenshot_2025-10-01_at_5.38.06_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222762c038-ca37-80ed-aefe-ce2183a9c641%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        },
        {
            'title': '🎨 Improved UI of Components View',
            'new': 'Enhanced the UI of the components view within the routine editor',
            'impact': 'Changed design of background, selected step, and step cards to increase usability by creating a more intuitive and aesthetic interface',
            'steps': [
                'Access the routine editor components view',
                'Experience improved visual design and layout',
                'Navigate components with enhanced usability'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3A1c2d8fb9-ff2b-4529-bb5e-4282309fd6ac%3AScreenshot_2025-10-01_at_5.39.14_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222772c038-ca37-802e-9350-f92e26765778%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        },
        {
            'title': '🔧 Consolidated Workflow Actions',
            'new': 'Consolidated workflow actions and zoom controls into a single floating pill on the left side of the screen',
            'impact': 'Improves the usability of the systems by combining platform related actions into a single area',
            'steps': [
                'Look for the floating pill on the left side of the screen',
                'Access workflow actions and zoom controls in one place',
                'Navigate platform features more efficiently'
            ],
            'file_urls': [
                '{{file://%7B%22source%22%3A%22attachment%3A6737f7d4-1d6e-4cc1-b382-2c45e07ebeec%3AScreenshot_2025-10-01_at_5.53.13_PM.png%22%2C%22permissionRecord%22%3A%7B%22table%22%3A%22block%22%2C%22id%22%3A%222782c038-ca37-80dc-aabd-ed46dab4b62c%22%2C%22spaceId%22%3A%2256486d48-de86-417a-ae69-1b2a5c1515db%22%7D%7D}}'
            ]
        }
    ]

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading('Routine Editor Newsletter', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date - properly bold formatted (NOT markdown)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime('%B %d, %Y'))
    date_run.bold = True

    # Add horizontal line
    doc.add_paragraph("---")

    # 1-Minute Summary
    doc.add_heading('1-Minute Summary', level=2)
    summary_items = [
        "Enhanced cross-system search with @ functionality makes finding content faster and more intuitive",
        "New evaluation workflow improvements streamline testing with direct eval-to-proposal linking",
        "Improved UI components and navigation across the platform create a more polished user experience"
    ]
    for item in summary_items:
        p = doc.add_paragraph()
        p.add_run(f"- {item}")

    doc.add_paragraph("---")

    # Feature Updates
    doc.add_heading('Feature Updates', level=2)

    images_embedded = 0

    for item in items:
        # Feature heading
        heading = doc.add_heading(item['title'], level=3)

        # Find and embed images for this feature
        image_paths = []
        for file_url in item['file_urls']:
            filename = extract_image_filename(file_url)
            if filename:
                local_path = find_matching_image(filename, snapshot_folder)
                if local_path:
                    image_paths.append(local_path)

        # Add images based on count
        if len(image_paths) == 1:
            # Single image - 6 inches wide, centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(image_paths[0], width=Inches(6))
            print(f"✅ Embedded 1 image for {item['title']}")
            images_embedded += 1

        elif len(image_paths) == 2:
            # Two images - 6 inches wide each, centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(image_paths[0], width=Inches(6))
            doc.add_picture(image_paths[1], width=Inches(6))
            print(f"✅ Embedded 2 images for {item['title']}")
            images_embedded += 2

        elif len(image_paths) > 2:
            # More than 2 images - add all as 6-inch centered
            for img_path in image_paths:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(img_path, width=Inches(6))
            print(f"✅ Embedded {len(image_paths)} images for {item['title']}")
            images_embedded += len(image_paths)

        # Feature content
        p = doc.add_paragraph()
        p.add_run("What's new: ").bold = True
        p.add_run(item['new'])

        p = doc.add_paragraph()
        p.add_run("Impact: ").bold = True
        p.add_run(item['impact'])

        p = doc.add_paragraph()
        p.add_run("How it works:").bold = True

        for i, step in enumerate(item['steps'], 1):
            p = doc.add_paragraph()
            p.add_run(f"\t{i}. {step}")

    doc.add_paragraph("---")

    # Behind the Scenes
    doc.add_heading('Behind the Scenes', level=2)

    doc.add_heading('Evaluation System Improvements', level=3)
    p = doc.add_paragraph()
    p.add_run("Impact: ").bold = True
    p.add_run("Faster evaluation creation and more consistent testing workflows")

    p = doc.add_paragraph()
    p.add_run("Changes:").bold = True
    changes = [
        "Evals from Simulations: Quickly create new eval cases and load existing eval cases as simulations",
        "Adding caching to simulations: Faster, cheaper, and more consistent LLMJ output"
    ]
    for change in changes:
        p = doc.add_paragraph()
        p.add_run(f"\t- {change}")

    doc.add_heading('Cross-Page Workflow Enhancement', level=3)
    p = doc.add_paragraph()
    p.add_run("Impact: ").bold = True
    p.add_run("Improved context visibility and workflow continuity")

    p = doc.add_paragraph()
    p.add_run("Changes:").bold = True
    p = doc.add_paragraph()
    p.add_run("\t- Enables cross-page workflow editing: Improves UI/UX by displaying active context as pills in Distillery Chat")

    doc.add_heading('User Experience Infrastructure', level=3)
    p = doc.add_paragraph()
    p.add_run("Impact: ").bold = True
    p.add_run("Better feedback collection and streamlined development processes")

    p = doc.add_paragraph()
    p.add_run("Changes:").bold = True
    changes = [
        "In-app feedback mechanisms (BE): Identifies opportunities to improve UI/UX of Distillery",
        "Migrates Tower Chat UI components to shadcn/ui: Standardizes design system components"
    ]
    for change in changes:
        p = doc.add_paragraph()
        p.add_run(f"\t- {change}")

    doc.add_paragraph("---")

    # Feedback & Support
    doc.add_heading('Feedback & Support', level=2)
    p = doc.add_paragraph("We'd love your input!")
    p = doc.add_paragraph()
    p.add_run("Feedback: ").bold = True
    p.add_run("Juan@distyl.ai")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Best regards,").bold = True
    p = doc.add_paragraph()
    p.add_run("The Distyl Team").italic = True

    # Save document
    doc.save('/Users/jcastillo/Desktop/Routine_Editor_Newsletter.docx')

    print(f"\n✅ Newsletter created successfully!")
    print(f"📄 Saved as: Routine_Editor_Newsletter.docx")
    print(f"🖼️  Embedded {images_embedded} real images from local snapshot folder")

    if images_embedded == 0:
        print("⚠️  No images were found in the snapshot folder")
        print(f"📁 Expected filenames:")
        for item in items:
            for file_url in item['file_urls']:
                filename = extract_image_filename(file_url)
                if filename:
                    print(f"   - {filename}")

if __name__ == "__main__":
    # Check if snapshot folder exists
    snapshot_folder = "/Users/jcastillo/Desktop/snapshot"

    print("🚀 Creating newsletter with local images...")
    print(f"📁 Looking for images in: {snapshot_folder}")

    if os.path.exists(snapshot_folder):
        files = os.listdir(snapshot_folder)
        print(f"🖼️  Found {len(files)} files in snapshot folder")
        for f in files[:5]:  # Show first 5 files
            print(f"   - {f}")
        if len(files) > 5:
            print(f"   ... and {len(files) - 5} more")
    else:
        print(f"❌ Snapshot folder not found: {snapshot_folder}")
        print("📝 Create the folder and add your screenshot files:")
        print(f"   mkdir -p {snapshot_folder}")
        print("   # Copy your screenshot files there")

    create_newsletter_with_local_images(snapshot_folder)